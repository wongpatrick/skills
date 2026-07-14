# /// script
# requires-python = ">=3.9"
# dependencies = [
#     "pypdf",
#     "python-docx",
# ]
# ///

import os
import sys
import argparse
from pypdf import PdfReader
from docx import Document

def extract_pdf(file_path):
    try:
        reader = PdfReader(file_path)
        text_parts = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text()
            if text:
                text_parts.append(text)
        return "\n\n--- Page Break ---\n\n".join(text_parts)
    except Exception as e:
        sys.stderr.write(f"Error reading PDF {file_path}: {e}\n")
        sys.exit(1)

def extract_docx(file_path):
    try:
        doc = Document(file_path)
        text_parts = []
        
        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
                
        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    text_parts.append(" | ".join(row_text))
                    
        return "\n".join(text_parts)
    except Exception as e:
        sys.stderr.write(f"Error reading DOCX {file_path}: {e}\n")
        sys.exit(1)

def extract_txt(file_path):
    try:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    except Exception as e:
        sys.stderr.write(f"Error reading text file {file_path}: {e}\n")
        sys.exit(1)

def main():
    parser = argparse.ArgumentParser(description="Extract text from PDF, DOCX, TXT, or MD files.")
    parser.add_argument("input", help="Path to the input file.")
    parser.add_argument("--output", required=True, help="Path to write the extracted text to.")
    
    args = parser.parse_args()
    
    if not os.path.exists(args.input):
        sys.stderr.write(f"Input file does not exist: {args.input}\n")
        sys.exit(1)
        
    ext = os.path.splitext(args.input.lower())[1]
    
    if ext == ".pdf":
        text = extract_pdf(args.input)
    elif ext == ".docx":
        text = extract_docx(args.input)
    elif ext in [".txt", ".md"]:
        text = extract_txt(args.input)
    else:
        sys.stderr.write(f"Unsupported file format: {ext}. Supported formats are .pdf, .docx, .txt, .md\n")
        sys.exit(1)
        
    try:
        # Create parent directories for output if they don't exist
        out_dir = os.path.dirname(args.output)
        if out_dir and not os.path.exists(out_dir):
            os.makedirs(out_dir, exist_ok=True)
            
        with open(args.output, "w", encoding="utf-8") as f:
            f.write(text)
        print(f"Success: Extracted text written to {args.output}")
    except Exception as e:
        sys.stderr.write(f"Error writing output to {args.output}: {e}\n")
        sys.exit(1)

if __name__ == "__main__":
    main()
